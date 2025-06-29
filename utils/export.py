from fpdf import FPDF
from docx import Document

def export_to_pdf(title: str, content: str, filename: str = "output.pdf") -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=title, align='L')
    pdf.ln()

    for paragraph in content.split("\n"):
        pdf.multi_cell(0, 10, txt=paragraph, align='L')

    pdf_output = bytes(pdf.output(dest='S').encode('latin1'))
    return pdf_output

def export_to_docx(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in content.split("\n"):
        doc.add_paragraph(paragraph)

    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
